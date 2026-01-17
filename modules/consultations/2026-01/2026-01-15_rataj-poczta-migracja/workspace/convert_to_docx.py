# Wygenerowano: 15.01.2026
# Skrypt konwertujący propozycję do formatu Word

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Ustawia kolor tła komórki tabeli"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(doc, headers, rows, header_color="D9E2F3"):
    """Dodaje tabelę do dokumentu"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Nagłówki
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], header_color)

    # Dane
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = str(cell_data)

    return table

doc = Document()

# Styl dokumentu
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# === NAGŁÓWEK ===
title = doc.add_heading('Propozycja rozwiązania: Archiwizacja poczty e-mail', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Metadane
p = doc.add_paragraph()
p.add_run('Dla: ').bold = True
p.add_run('Rataj Sp. z o.o.\n')
p.add_run('Data: ').bold = True
p.add_run('15.01.2026\n')
p.add_run('Dotyczy: ').bold = True
p.add_run('Poczta archiwalna z Windows Live Mail')

doc.add_paragraph()

# === SYTUACJA WYJŚCIOWA ===
doc.add_heading('Sytuacja wyjściowa', level=1)

p = doc.add_paragraph()
p.add_run('Obecnie poczta firmowa przechowywana jest w programie ')
p.add_run('Windows Live Mail').bold = True
p.add_run(', który:')

doc.add_paragraph('Nie jest już wspierany przez Microsoft', style='List Bullet')
doc.add_paragraph('Działa, ale z ryzykiem awarii przy dużych wolumenach danych', style='List Bullet')
doc.add_paragraph('Przechowuje pocztę w formacie EML (każdy e-mail = osobny plik)', style='List Bullet')

doc.add_heading('Dane techniczne', level=2)
add_table(doc,
    ['Parametr', 'Wartość'],
    [
        ['Rozmiar poczty', '40–70 GB'],
        ['Hosting (CyberFolks)', 'Pakiet 100 GB, wykorzystane ~22 GB'],
        ['Wolne miejsce na serwerze', '~78 GB (wystarczy na wgranie archiwum)'],
    ])

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Cel: ').bold = True
p.add_run('Zabezpieczenie poczty archiwalnej i zapewnienie wygodnego dostępu do niej na przyszłość.')

# === ZNANE RYZYKA ===
doc.add_heading('Znane ryzyka migracji (dotyczy wszystkich wariantów)', level=2)

p = doc.add_paragraph()
p.add_run('⚠️ UWAGA: ').bold = True
p.add_run('W obecnym archiwum mogą znajdować się:')

add_table(doc,
    ['Problem', 'Opis'],
    [
        ['Uszkodzone wiadomości', 'Po latach użytkowania i kilku awariach bazy WLM część wiadomości może być uszkodzona (nieczytelna, niepełna).'],
        ['Zbyt duże wiadomości', 'Wiadomości z załącznikami >25-50 MB mogą przekraczać limity obecnych serwerów pocztowych.'],
        ['Niekompatybilne formaty', 'Stare wiadomości mogą zawierać elementy nieobsługiwane przez nowe systemy.'],
    ])

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Co to oznacza:').bold = True

doc.add_paragraph('Część wiadomości może nie dać się zmigrować automatycznie', style='List Bullet')
doc.add_paragraph('Wymagane będzie ręczne rozwiązanie dla problematycznych przypadków', style='List Bullet')
doc.add_paragraph('Na etapie wstępnym nie da się dokładnie określić, ile takich wiadomości jest', style='List Bullet')
doc.add_paragraph('Dopiero podczas migracji testowej wyjdą konkretne problemy', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Zalecenie: ').bold = True
p.add_run('Przed pełną migracją wykonać test na próbce (np. 1 folder) i zidentyfikować typowe problemy.')

# === WARIANT 1 ===
doc.add_page_break()
doc.add_heading('Wariant 1: Thunderbird z folderami lokalnymi', level=1)

doc.add_heading('Na czym polega?', level=2)
p = doc.add_paragraph()
p.add_run('Poczta archiwalna zostaje przeniesiona do programu ')
p.add_run('Mozilla Thunderbird').bold = True
p.add_run(' — nowoczesnego, bezpłatnego klienta pocztowego. Wiadomości trafiają do ')
p.add_run('folderów lokalnych').bold = True
p.add_run(' (przechowywanych na dysku komputera, nie na serwerze).')

doc.add_heading('Zalety', level=2)
add_table(doc,
    ['Zaleta', 'Opis'],
    [
        ['Bezpłatne', 'Thunderbird jest darmowy'],
        ['Jeden program', 'Archiwum i bieżąca poczta w jednym miejscu'],
        ['Wyszukiwanie', 'Szybkie przeszukiwanie wszystkich wiadomości'],
        ['Aktywny rozwój', 'Program jest rozwijany (nowe wersje co miesiąc)'],
        ['Prosta migracja komputera', 'Kopiowanie jednego folderu na nowy komputer'],
        ['Dane pozostają lokalne', 'Poczta nie wraca do "chmury"'],
    ])

doc.add_heading('Wady i ryzyka', level=2)
add_table(doc,
    ['Ryzyko', 'Wyjaśnienie'],
    [
        ['Format MBOX', 'Thunderbird przechowuje wszystkie e-maile z jednego folderu w jednym dużym pliku.'],
        ['Uszkodzenie = utrata folderu', 'Jeśli plik się uszkodzi, można stracić cały folder naraz.'],
        ['Duże foldery = wolniejsze', 'Foldery powyżej 2-4 GB mogą działać wolniej.'],
        ['Brak dostępu mobilnego', 'Archiwum dostępne tylko z tego komputera.'],
    ])

doc.add_heading('Jak minimalizować ryzyko?', level=2)
doc.add_paragraph('Podział dużych folderów — np. Inbox podzielić na Inbox_2020, Inbox_2021, itd.', style='List Number')
doc.add_paragraph('Regularne backupy — kopia folderu Thunderbird na zewnętrzny dysk (co tydzień)', style='List Number')
doc.add_paragraph('Zachowanie oryginalnych plików EML — jako dodatkowa kopia bezpieczeństwa', style='List Number')

doc.add_heading('Koszty', level=2)
add_table(doc,
    ['Pozycja', 'Koszt'],
    [
        ['Program', '0 zł'],
        ['Wdrożenie*', 'Do ustalenia'],
        ['Utrzymanie roczne', '0 zł'],
    ])
doc.add_paragraph('*Import danych, konfiguracja, szkolenie')

# === WARIANT 2 ===
doc.add_page_break()
doc.add_heading('Wariant 2: MailStore', level=1)

p = doc.add_paragraph()
p.add_run('MailStore').bold = True
p.add_run(' to specjalistyczne oprogramowanie do archiwizacji poczty e-mail. Tworzy bezpieczne, przeszukiwalne archiwum wszystkich wiadomości.')

doc.add_heading('Opcja 2A: MailStore Home (jeden komputer)', level=2)

p = doc.add_paragraph()
p.add_run('Uwaga: ').bold = True
p.add_run('MailStore Home jest bezpłatny, ale tylko do użytku prywatnego. Dla firmy formalnie wymagana jest wersja Server.')

add_table(doc,
    ['Zalety', 'Wady'],
    [
        ['Bezpłatne (ale tylko prywatnie!)', 'Dostęp tylko z jednego komputera'],
        ['Własny format (odporny na błędy)', 'Brak interfejsu webowego'],
        ['Szybkie wyszukiwanie', 'Nie do użytku firmowego (licencja)'],
        ['Zachowuje strukturę folderów', 'Brak dostępu mobilnego'],
        ['Dane pozostają lokalne', ''],
    ])

doc.add_heading('Opcja 2B: MailStore Server (dostęp sieciowy)', level=2)

add_table(doc,
    ['Zalety', 'Wady'],
    [
        ['Legalny do użytku firmowego', 'Płatne (patrz koszty poniżej)'],
        ['Dostęp z każdego komputera (web)', 'Wymaga komputera jako "serwera"'],
        ['Wielu użytkowników jednocześnie', 'Konfiguracja bardziej złożona'],
        ['Własny format (bezpieczny)', 'Wymaga dbania o lokalny serwer'],
        ['Wyszukiwanie w załącznikach', ''],
        ['Eksport z powrotem do plików EML', ''],
        ['Dane pozostają w firmie', ''],
    ])

doc.add_heading('Koszty MailStore Server (do 5 użytkowników)', level=2)
add_table(doc,
    ['Pozycja', 'Koszt EUR', 'Koszt PLN (ok.)'],
    [
        ['Licencja + 1 rok wsparcia', '€295', '~1 270 zł'],
        ['Odnowienie wsparcia (roczne)', '€85', '~365 zł/rok'],
    ])

p = doc.add_paragraph()
p.add_run('Co daje wsparcie?').bold = True
doc.add_paragraph('Aktualizacje programu', style='List Bullet')
doc.add_paragraph('Pomoc techniczna producenta', style='List Bullet')
doc.add_paragraph('Bez odnowienia: program działa dalej, ale bez nowych wersji', style='List Bullet')

# === WARIANT 3 ===
doc.add_page_break()
doc.add_heading('Wariant 3: Wgranie archiwum na serwer CyberFolks (IMAP)', level=1)

doc.add_heading('Na czym polega?', level=2)
p = doc.add_paragraph()
p.add_run('Poczta archiwalna zostaje wgrana z powrotem na serwer hostingowy CyberFolks. Dzięki temu jest dostępna przez IMAP z dowolnego urządzenia — komputera, telefonu, tabletu, a także przez webmail.')

doc.add_heading('Jak działa synchronizacja z Thunderbirdem?', level=2)
add_table(doc,
    ['Tryb', 'Jak działa'],
    [
        ['Online (domyślny)', 'Pobiera tylko nagłówki. Treść ładuje się po kliknięciu. Wymaga internetu.'],
        ['Offline', 'Pobiera całą pocztę na dysk lokalny. Można czytać bez internetu.'],
    ])

p = doc.add_paragraph()
p.add_run('Uwaga: ').bold = True
p.add_run('W trybie offline Thunderbird tworzy lokalną kopię w formacie MBOX. Przy 40-70 GB poczty, tyle samo miejsca zajmie na dysku komputera.')

doc.add_heading('Zalety', level=2)
add_table(doc,
    ['Zaleta', 'Opis'],
    [
        ['Dostęp z każdego urządzenia', 'Komputer, telefon, tablet, webmail'],
        ['Synchronizacja folderów', 'Zmiany widoczne wszędzie'],
        ['Backup po stronie serwera', 'Hosting robi kopie zapasowe'],
        ['Brak dodatkowych kosztów', 'Wykorzystanie istniejącego hostingu'],
    ])

doc.add_heading('Wady i ryzyka', level=2)
add_table(doc,
    ['Ryzyko', 'Wyjaśnienie'],
    [
        ['Poczta wraca do "chmury"', 'Dane na serwerze zewnętrznym (hosting)'],
        ['Czas wgrywania', 'Upload 40-70 GB może trwać wiele godzin/dni'],
        ['Limity załączników', 'Wiadomości >25-50 MB mogą być odrzucane'],
        ['Zużycie miejsca na hostingu', 'Po wgraniu zostanie tylko 8-38 GB wolnego'],
        ['Wolna pierwsza synchronizacja', 'Thunderbird będzie długo pobierał całość'],
        ['Zależność od dostawcy', 'Awaria hostingu = brak dostępu do archiwum'],
    ])

doc.add_heading('Koszty', level=2)
add_table(doc,
    ['Pozycja', 'Koszt'],
    [
        ['Dodatkowe oprogramowanie', '0 zł (wykorzystanie istniejącego hostingu)'],
        ['Wdrożenie*', 'Do ustalenia'],
        ['Utrzymanie roczne', '0 zł (w ramach hostingu)'],
    ])
doc.add_paragraph('*Eksport, upload, weryfikacja — przy 40-70 GB może być czasochłonne')

# === WARIANT 4 ===
doc.add_page_break()
doc.add_heading('Wariant 4: Migracja do Gmail z etykietami', level=1)

doc.add_heading('Na czym polega?', level=2)
p = doc.add_paragraph()
p.add_run('Cała poczta archiwalna zostaje przeniesiona na konto ')
p.add_run('Gmail / Google Workspace').bold = True
p.add_run('. Struktura folderów odwzorowana jest za pomocą ')
p.add_run('etykiet').bold = True
p.add_run(' (labels) — specyficznej funkcji Gmaila.')

p = doc.add_paragraph('Gmail może również:')
doc.add_paragraph('Pobierać nową pocztę z oryginalnej skrzynki ("Sprawdź pocztę z innych kont")', style='List Bullet')
doc.add_paragraph('Wysyłać w imieniu oryginalnego adresu ("Wyślij jako")', style='List Bullet')

doc.add_heading('Szczególne zalety Gmaila', level=2)
add_table(doc,
    ['Funkcja', 'Opis'],
    [
        ['Inteligentne wyszukiwanie', 'Bardzo szybkie, rozumie kontekst, szuka w załącznikach'],
        ['Deduplikacja', 'Jeśli ta sama wiadomość jest w wielu miejscach, Gmail zachowa jedną kopię'],
        ['Znany interfejs', 'Większość osób zna Gmaila'],
        ['Etykiety > foldery', 'Jedna wiadomość może mieć wiele etykiet'],
        ['Trwałość Google', 'Google zapewnia wysoką dostępność i backup'],
        ['Integracja', 'Kalendarz, Dysk, Meet — wszystko w jednym'],
    ])

doc.add_heading('Jak działa "wszystko w jednym"?', level=2)
p = doc.add_paragraph('Po skonfigurowaniu Gmail może:')
doc.add_paragraph('Pobierać pocztę z oryginalnej skrzynki (np. poczta@rataj.pl) automatycznie', style='List Number')
doc.add_paragraph('Wysyłać odpowiedzi tak, że odbiorca widzi adres @rataj.pl, nie @gmail.com', style='List Number')
doc.add_paragraph('Czyli: pracujesz w Gmailu, ale "na zewnątrz" wygląda to jak normalna poczta firmowa', style='List Number')

doc.add_heading('Wady i ryzyka', level=2)
add_table(doc,
    ['Ryzyko', 'Wyjaśnienie'],
    [
        ['Dane w chmurze Google', 'Poczta przechowywana na serwerach Google (USA)'],
        ['Wymaga subskrypcji', '40-70 GB przekracza darmowe 15 GB'],
        ['Zależność od Google', 'Zmiana polityki cenowej, warunków usługi'],
        ['Prywatność', 'Google analizuje treść (dla reklam w wersji darmowej)'],
        ['Migracja = jednorazowy wysiłek', 'Import dużej ilości danych może być skomplikowany'],
        ['Limity wiadomości', 'Gmail ma limit 25 MB na wiadomość'],
    ])

doc.add_heading('Koszty Google One / Google Workspace', level=2)
p = doc.add_paragraph()
p.add_run('Dla 40-70 GB poczty potrzebny jest płatny plan:')

add_table(doc,
    ['Plan', 'Pojemność', 'Koszt miesięczny', 'Koszt roczny (ok.)'],
    [
        ['Google One 100 GB', '100 GB', '~9 zł', '~108 zł'],
        ['Google One 200 GB', '200 GB', '~15 zł', '~180 zł'],
        ['Google Workspace Basic', '30 GB/user', '~27 zł/user', '~324 zł/user'],
    ])

p = doc.add_paragraph()
p.add_run('Uwaga: ').bold = True
p.add_run('Google One można współdzielić z rodziną (do 5 osób). Jeśli firma już ma pakiet Google One, może wystarczyć jego zwiększenie.')

# === PORÓWNANIE ===
doc.add_page_break()
doc.add_heading('Porównanie wszystkich wariantów', level=1)

# Tabela porównawcza - podzielona dla czytelności
add_table(doc,
    ['Aspekt', 'Thunderbird', 'MailStore Home', 'MailStore Server', 'IMAP CyberFolks', 'Gmail'],
    [
        ['Koszt początkowy', '0 zł', '0 zł', '~1 270 zł', '0 zł', '0 zł'],
        ['Koszt roczny', '0 zł', '0 zł', '~365 zł', '0 zł', '~108-180 zł'],
        ['Legalność dla firmy', '✓ Tak', '✗ Prywatnie', '✓ Tak', '✓ Tak', '✓ Tak'],
        ['Dostęp mobilny', '✗ Nie', '✗ Nie', '✓ Web (sieć lok.)', '✓ Tak', '✓ Tak'],
        ['Dostęp z internetu', '✗ Nie', '✗ Nie', '✗ Sieć lokalna', '✓ Tak', '✓ Tak'],
        ['Dane w chmurze', '✗ Lokalne', '✗ Lokalne', '✗ Lokalne', '! Hosting', '! Google'],
        ['Wyszukiwanie', '✓ Dobre', '✓ Bardzo dobre', '✓ Bardzo dobre', '✓ Dobre', '✓ Najlepsze'],
        ['Deduplikacja', '✗ Nie', '✓ Tak', '✓ Tak', '✗ Nie', '✓ Tak'],
        ['Bezpieczeństwo formatu', '! MBOX', '✓ Własny', '✓ Własny', '! Serwer', '✓ Google'],
        ['Łatwość wdrożenia', '✓ Prosta', '✓ Prosta', '! Średnia', '! Czasochłonne', '! Średnia'],
        ['Integracja z telefonem', '✗ Nie', '✗ Nie', '! Tylko web', '✓ Tak', '✓ Najlepsza'],
        ['Tolerancja dużych/uszkodzonych', '✓ Wysoka', '✓ Wysoka', '✓ Wysoka', '! Limity serwera', '! Limit 25 MB'],
    ])

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Legenda: ').bold = True
p.add_run('✓ = zaleta/tak, ✗ = wada/nie, ! = uwaga/ograniczenie')

# === REKOMENDACJE ===
doc.add_heading('Rekomendacje według potrzeb', level=1)

add_table(doc,
    ['Jeśli priorytetem jest...', '...to zalecamy'],
    [
        ['Minimalne koszty + dane lokalne', 'Wariant 1: Thunderbird + backup'],
        ['Profesjonalne archiwum w firmie', 'Wariant 2B: MailStore Server'],
        ['Dostęp mobilny + obecna infrastruktura', 'Wariant 3: IMAP CyberFolks'],
        ['Najlepsze wyszukiwanie + wygoda', 'Wariant 4: Gmail'],
        ['Dane muszą zostać w firmie (poufność)', 'Wariant 1 lub 2B (nie chmura)'],
    ])

# === NASTĘPNE KROKI ===
doc.add_heading('Następne kroki', level=1)

p = doc.add_paragraph('Po wyborze wariantu:')
doc.add_paragraph('Wykonanie pełnego backupu obecnej poczty (40-70 GB)', style='List Number')
doc.add_paragraph('Sprawdzenie dokładnego rozmiaru i struktury folderów', style='List Number')
doc.add_paragraph('Identyfikacja wiadomości >25 MB (potencjalne problemy)', style='List Number')
doc.add_paragraph('Test migracji na małej próbce (np. 1 folder, 1000 wiadomości)', style='List Number')
doc.add_paragraph('Pełna migracja', style='List Number')
doc.add_paragraph('Weryfikacja poprawności przeniesienia', style='List Number')
doc.add_paragraph('Szkolenie z obsługi nowego rozwiązania', style='List Number')

# === PYTANIA ===
doc.add_heading('Pytania do podjęcia decyzji', level=1)

doc.add_paragraph('Czy dostęp mobilny (telefon, przeglądarka) jest ważny?', style='List Number')
doc.add_paragraph('Czy poczta może być przechowywana w chmurze (hosting/Google)?', style='List Number')
doc.add_paragraph('Jaki jest budżet na rozwiązanie?', style='List Number')
doc.add_paragraph('Ile osób potrzebuje dostępu do archiwum?', style='List Number')
doc.add_paragraph('Czy istotna jest deduplikacja (usunięcie powtórzonych wiadomości)?', style='List Number')

# === STOPKA ===
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Dokument roboczy — wersja 2.1').italic = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Zapis
output_path = r"c:\WORK\projects\_tools-and-lab\consultations\2026-01\2026-01-15_rataj-poczta-migracja\workspace\PROPOZYCJA_DLA_KLIENTA.docx"
doc.save(output_path)
print(f"Dokument zapisany: {output_path}")
